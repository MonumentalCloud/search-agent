#!/usr/bin/env python3
"""
Script to ingest meeting logs from data folder into the existing fiqa_chroma database.
This creates a combined database with both FiQA financial data and meeting logs.
"""

import os
import sys
import logging
import json
import docx
from pathlib import Path
from datetime import datetime
from typing import Dict, Any

# Add project root to path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

from adapters.chroma_adapter import ChromaClient
from configs.load import get_default_llm, get_default_embeddings

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/meetings_to_fiqa_ingestion.log'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from DOCX including tables."""
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                full_text.append(' | '.join(row_text))
    
    return '\n'.join(full_text)

def llm_extract_all_metadata(text: str) -> dict:
    """Extract all metadata using a single LLM pass."""
    llm = get_default_llm()
    
    prompt = f"""
You are an expert at extracting key information from meeting minutes documents.
Analyze the following meeting minutes and extract the following details:
1. The actual meeting date (primary_date) in YYYY-MM-DD format.
2. A list of all attendees (attendees).
3. A list of key decisions (key_decisions).
4. A list of action items (action_items).
5. The meeting location (location).
6. The meeting topic (topic).

If a date is mentioned without a year, assume the current year (2025).
Return your answer as a JSON object with the following keys:
{{
  "primary_date": "YYYY-MM-DD",
  "attendees": ["Name1", "Name2"],
  "key_decisions": ["Decision 1", "Decision 2"],
  "action_items": ["Action 1", "Action 2"],
  "location": "Meeting Room Name",
  "topic": "Meeting Topic"
}}

Text:
{text}
"""
    
    try:
        response = llm.invoke(prompt)
        # Parse JSON from response
        import re
        json_match = re.search(r'\{.*\}', response.content, re.DOTALL)
        if json_match:
            metadata = json.loads(json_match.group())
            # Rename primary_date to meeting_date for consistency
            if 'primary_date' in metadata:
                metadata['meeting_date'] = metadata.pop('primary_date')
            return metadata
        else:
            logger.error("No JSON found in LLM response")
            return {}
    except Exception as e:
        logger.error(f"LLM extraction failed: {e}")
        return {}

def ingest_meeting_to_fiqa(file_path: str, client: ChromaClient, embeddings_model) -> Dict[str, Any]:
    """
    Ingest a single meeting document into the fiqa_chroma database.
    
    Args:
        file_path: Path to the meeting document file.
        client: ChromaClient instance.
        embeddings_model: The embedding model to use.
        
    Returns:
        Dictionary with ingestion results and extracted metadata
    """
    logger.info(f"📄 Processing: {Path(file_path).name}")
    
    try:
        # Extract text from document
        text = extract_text_from_docx(file_path)
        logger.info(f"✅ Extracted {len(text)} characters")
        
        # Extract metadata with single LLM pass
        logger.info("🤖 Running single LLM pass for metadata extraction...")
        metadata = llm_extract_all_metadata(text)
        
        if not metadata:
            logger.error("❌ Failed to extract metadata from LLM")
            return {"success": False, "file": Path(file_path).name, "reason": "LLM metadata extraction failed"}
        
        # Prepare metadata for ChromaDB
        doc_id = Path(file_path).stem
        chroma_metadata = {
            "file_name": Path(file_path).name,
            "file_path": str(file_path),
            "ingestion_date": datetime.now().isoformat(),
            "doc_id": doc_id,
            "summary": metadata.get("topic", "No topic extracted"),
            "meeting_date": metadata.get("meeting_date"),
            "valid_from": metadata.get("meeting_date"), # Use meeting_date as valid_from
            "attendees": json.dumps(metadata.get("attendees", [])),
            "key_decisions": json.dumps(metadata.get("key_decisions", [])),
            "action_items": json.dumps(metadata.get("action_items", [])),
            "location": metadata.get("location", ""),
            "topic": metadata.get("topic", ""),
            "all_dates": json.dumps([metadata.get("meeting_date")] if metadata.get("meeting_date") else []),
            # Add source info to distinguish from FiQA data
            "source_type": "meeting_log",
            "document_type": "meeting_minutes"
        }
        
        # Filter out None values from metadata
        chroma_metadata = {k: v for k, v in chroma_metadata.items() if v is not None}
        
        # Generate embedding for the entire document
        document_embedding = embeddings_model.embed_query(text)
        
        # Add to the fiqa_full collection (same collection as FiQA data)
        client._client.get_collection("fiqa_full").add(
            ids=[doc_id],
            documents=[text],
            metadatas=[chroma_metadata],
            embeddings=[document_embedding]
        )
        
        logger.info("✅ Metadata extracted successfully!")
        logger.info(f"📅 Meeting date: {metadata.get('meeting_date', 'N/A')}")
        logger.info(f"👥 Attendees: {len(metadata.get('attendees', []))}")
        logger.info(f"📋 Decisions: {len(metadata.get('key_decisions', []))}")
        logger.info(f"✅ Actions: {len(metadata.get('action_items', []))}")
        logger.info(f"✅ Successfully ingested {Path(file_path).name} into fiqa_full collection")
        
        return {"success": True, "file": Path(file_path).name, "metadata": metadata}
    
    except Exception as e:
        logger.error(f"❌ Error ingesting {Path(file_path).name}: {e}", exc_info=True)
        return {"success": False, "file": Path(file_path).name, "reason": str(e)}

def main():
    """Main function to ingest meeting logs into fiqa_chroma."""
    
    logger.info("🎯 Ingesting Meeting Logs into FiQA Database")
    logger.info("=" * 60)
    
    # Initialize Chroma client and embeddings model
    # Temporarily switch to fiqa_chroma database
    logger.info("🔄 Switching to fiqa_chroma database for ingestion...")
    
    # Create a temporary client that points to fiqa_chroma
    import chromadb
    fiqa_client = chromadb.PersistentClient(path="./fiqa_chroma")
    
    # Create a custom ChromaClient that uses the fiqa_client
    class FiQAChromaClient:
        def __init__(self):
            self._client = fiqa_client
            self._connected = True
            self.chunk_collection = "fiqa_full"
            self.document_collection = "Document"
    
    client = FiQAChromaClient()
    logger.info("✅ Connected to fiqa_chroma database")
    
    embeddings_model = get_default_embeddings()
    
    data_dir = project_root / "data"
    docx_files = list(data_dir.glob("*.docx"))
    
    if not docx_files:
        logger.warning("No DOCX files found in the 'data' directory.")
        return
    
    logger.info(f"📁 Found {len(docx_files)} DOCX files to process")
    
    successful_ingestions = []
    failed_ingestions = []
    extracted_dates = {}
    
    for docx_file in sorted(docx_files):
        result = ingest_meeting_to_fiqa(str(docx_file), client, embeddings_model)
        if result["success"]:
            successful_ingestions.append(result)
            if result["metadata"] and "meeting_date" in result["metadata"]:
                extracted_dates[result["file"]] = result["metadata"]["meeting_date"]
        else:
            failed_ingestions.append(result)
    
    logger.info("=" * 60)
    logger.info("📊 INGESTION SUMMARY")
    logger.info("=" * 60)
    logger.info(f"✅ Successful: {len(successful_ingestions)}")
    logger.info(f"❌ Failed: {len(failed_ingestions)}")
    logger.info(f"📁 Total files: {len(docx_files)}")
    logger.info("\n📅 EXTRACTED MEETING DATES:")
    for file_name, date in extracted_dates.items():
        logger.info(f"  - {file_name}: {date}")
    
    if failed_ingestions:
        logger.error("❌ Some ingestions failed")
        for fail in failed_ingestions:
            logger.error(f"  - {fail['file']}: {fail.get('reason', 'Unknown error')}")
    else:
        logger.info("🎉 All meeting logs successfully ingested into fiqa_chroma!")
        logger.info("💡 The fiqa_chroma database now contains both FiQA data and meeting logs")
        logger.info("💡 You can test this by switching to fiqa database and asking about meetings")

if __name__ == "__main__":
    main()
