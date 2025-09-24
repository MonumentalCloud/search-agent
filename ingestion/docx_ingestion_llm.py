"""
Module for ingesting DOCX files with LLM-based date extraction.
"""
import os
import json
import re
from pathlib import Path
from typing import Dict, List, Optional, Union, Any
import datetime
import logging
import docx
from docx import Document

from ingestion.hybrid_chunking import hybrid_chunk_text
from adapters.weaviate_adapter import WeaviateClient
from configs.load import get_default_llm

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = Document(file_path)
        full_text = []
        
        # Extract metadata from document properties if available
        metadata = {}
        try:
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
        except Exception as e:
            logger.warning(f"Could not extract metadata from {file_path}: {e}")
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        # Add metadata as a header if available
        header_lines = []
        if "title" in metadata:
            header_lines.append(f"# {metadata['title']}")
        if "created" in metadata:
            created_date = metadata["created"].split("T")[0] if "T" in metadata["created"] else metadata["created"]
            header_lines.append(f"Date: {created_date}")
        if "author" in metadata:
            header_lines.append(f"Author: {metadata['author']}")
        if "subject" in metadata:
            header_lines.append(f"Subject: {metadata['subject']}")
        if "keywords" in metadata:
            header_lines.append(f"Keywords: {metadata['keywords']}")
        
        # Combine header and text
        if header_lines:
            full_text = header_lines + [""] + full_text
        
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX file {file_path}: {e}")
        return ""

def extract_metadata_from_filename(file_path: str) -> Dict[str, str]:
    """Extract metadata from Korean meeting filename like '회의록_01_마케팅.docx'."""
    try:
        filename = os.path.basename(file_path)
        name_parts = filename.split('_')
        
        metadata = {
            "doc_type": "회의록" if len(name_parts) > 0 and name_parts[0] == "회의록" else "document",
        }
        
        # Extract meeting number if available
        if len(name_parts) > 1 and name_parts[1].isdigit():
            metadata["meeting_number"] = name_parts[1]
        
        # Extract meeting topic if available
        if len(name_parts) > 2:
            topic = name_parts[2].split('.')[0]  # Remove file extension
            metadata["topic"] = topic
        
        return metadata
    except Exception as e:
        logger.warning(f"Could not extract metadata from filename {file_path}: {e}")
        return {"doc_type": "document"}

def extract_dates_with_llm(text: str) -> Dict[str, Any]:
    """
    Extract dates from text using LLM.
    
    Args:
        text: The text to extract dates from
        
    Returns:
        Dictionary with extracted dates and their context
    """
    llm = get_default_llm()
    
    # Create the prompt for the LLM
    prompt = f"""
You are an expert at analyzing text and extracting date information. 
Extract ALL dates mentioned in the following text, paying special attention to meeting dates, event dates, and document dates.

IMPORTANT RULES:
1. Look for dates in various formats (YYYY-MM-DD, MM/DD/YYYY, Korean format like YYYY년 MM월 DD일, etc.)
2. Pay special attention to phrases like "일시", "날짜", "Date:", "다음 회의:", etc. that indicate dates
3. For each date found, identify its context (e.g., "meeting date", "document date", "next meeting", etc.)
4. Always normalize dates to ISO format (YYYY-MM-DD)
5. The PRIMARY date should be the actual meeting/event date, not the "next meeting" date
6. Look for dates associated with "일시" (meaning "date/time") as these often indicate the actual meeting date

Return a JSON object with the following structure:
{{
  "primary_date": "YYYY-MM-DD",  // The actual meeting/event date (not "next meeting" date)
  "primary_date_context": "brief explanation of what this date represents",
  "all_dates": [
    {{
      "date": "YYYY-MM-DD",
      "context": "explanation of what this date represents",
      "original_text": "the original text snippet containing the date"
    }},
    ...
  ]
}}

Text:
{text[:2000]}  // Truncate to first 2000 chars for LLM context window
"""
    
    try:
        # Fix for langchain deprecation warning - use invoke() instead of __call__
        response = llm.invoke(prompt)
        
        # Handle different response types based on LLM provider
        response_text = ""
        if hasattr(response, "content"):
            # For ChatCompletion style responses
            response_text = response.content
        elif isinstance(response, str):
            # For string responses
            response_text = response
        else:
            # Try to extract content from other response types
            try:
                response_text = str(response)
            except:
                logger.error(f"Unexpected response type: {type(response)}")
                return {"primary_date": None, "all_dates": []}
        
        # Parse the JSON response
        try:
            # Find JSON in the response
            json_match = re.search(r'```json\n(.*?)\n```', response_text, re.DOTALL)
            if json_match:
                json_str = json_match.group(1)
            else:
                # Try to find JSON without markdown formatting
                json_match = re.search(r'({.*})', response_text, re.DOTALL)
                if json_match:
                    json_str = json_match.group(1)
                else:
                    json_str = response_text
            
            # Clean up the JSON string
            json_str = json_str.strip()
            if json_str.startswith('```') and json_str.endswith('```'):
                json_str = json_str[3:-3].strip()
                
            result = json.loads(json_str)
            return result
        except json.JSONDecodeError as e:
            logger.error(f"Error parsing LLM response: {e}")
            logger.debug(f"Response: {response_text}")
            return {"primary_date": None, "all_dates": []}
            
    except Exception as e:
        logger.error(f"Error extracting dates with LLM: {e}")
        return {"primary_date": None, "all_dates": []}

def ingest_docx(file_path: str, client: WeaviateClient) -> bool:
    """Ingest a DOCX file into Weaviate."""
    try:
        # Extract text from DOCX
        text = extract_text_from_docx(file_path)
        if not text:
            logger.warning(f"No text extracted from {file_path}")
            return False
        
        # Extract metadata from filename
        metadata = extract_metadata_from_filename(file_path)
        
        # Extract dates using LLM
        date_info = extract_dates_with_llm(text)
        meeting_date = date_info.get("primary_date")
        
        if meeting_date:
            metadata["meeting_date"] = meeting_date
            metadata["date_context"] = date_info.get("primary_date_context", "")
            
            # Store all extracted dates in metadata
            all_dates = []
            for date_entry in date_info.get("all_dates", []):
                all_dates.append({
                    "date": date_entry.get("date"),
                    "context": date_entry.get("context"),
                    "original_text": date_entry.get("original_text")
                })
            metadata["all_dates"] = all_dates
            
            print(f"Extracted primary date: {meeting_date} ({metadata.get('date_context')}) from {file_path}")
            print(f"Found {len(all_dates)} dates in the document")
        
        # Generate a unique document ID
        doc_id = f"docx_{Path(file_path).stem}"
        
        # Create document object
        document = {
            "doc_id": doc_id,
            "title": metadata.get("topic", Path(file_path).stem),
            "doc_type": metadata.get("doc_type", "document"),
            "created_at": datetime.datetime.now().isoformat(),
            "updated_at": datetime.datetime.now().isoformat(),
            "metadata": metadata,
        }
        
        # Insert document into Weaviate
        client.batch_upsert_documents([document])
        
        # Chunk the text using hybrid chunking
        chunks = hybrid_chunk_text(text)
        
        # Prepare chunks for Weaviate
        weaviate_chunks = []
        for i, chunk in enumerate(chunks):
            chunk_id = f"{doc_id}_{i}"
            
            # Create chunk with metadata
            weaviate_chunk = {
                "chunk_id": chunk_id,
                "doc_id": doc_id,
                "section": chunk.get("summary") or chunk.get("section_title") or f"Section {i+1}",
                "body": chunk.get("text") or chunk.get("body") or "",
                "entities": chunk.get("entities") or [],
                "relationships": chunk.get("relationships") or [],
                "created_at": datetime.datetime.now().isoformat(),
                "updated_at": datetime.datetime.now().isoformat(),
            }
            
            # Add meeting date as valid_from if available
            if meeting_date:
                # Format as RFC3339 with time and timezone for Weaviate
                # Don't add Z here since WeaviateClient will add it
                formatted_date = f"{meeting_date}T00:00:00"
                weaviate_chunk["valid_from"] = formatted_date
                print(f"Setting valid_from to {formatted_date} for chunk {chunk_id}")
            
            weaviate_chunks.append(weaviate_chunk)
        
        # Insert chunks into Weaviate
        if weaviate_chunks:
            client.batch_upsert_chunks(weaviate_chunks)
            logger.info(f"Ingested {file_path} with {len(weaviate_chunks)} chunks")
            return True
        else:
            logger.warning(f"No chunks created for {file_path}")
            return False
        
    except Exception as e:
        logger.error(f"Error ingesting DOCX file {file_path}: {e}")
        return False

def ingest_docx_directory(directory: str, client: WeaviateClient, file_pattern: str = "*.docx") -> Dict[str, bool]:
    """Ingest all DOCX files in a directory into Weaviate."""
    results = {}
    directory_path = Path(directory)
    
    # Find all DOCX files in the directory
    docx_files = list(directory_path.glob(file_pattern))
    logger.info(f"Found {len(docx_files)} DOCX files in {directory}")
    
    # Ingest each file
    for file_path in docx_files:
        logger.info(f"Ingesting {file_path}")
        success = ingest_docx(str(file_path), client)
        results[str(file_path)] = success
    
    return results