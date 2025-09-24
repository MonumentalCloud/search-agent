"""
Module for processing documents with LLM-based date extraction, adapted to GenOS format.
This is a self-contained module that can be directly integrated into GenOS.
"""
import os
import json
import re
import uuid
import logging
import docx
from datetime import datetime
from typing import Dict, List, Any, Optional
from pathlib import Path
import re

# Use requests library for GenOS API
import json
import os
import requests
import time

logger = logging.getLogger(__name__)


class GenosClient:
    """GenOS API client using requests library."""
    
    def __init__(self, api_key: str = "85a19cbab1454db29d0b1460304d679e", model: str = "qwen2.5 72b instruct"):
        """Initialize with API key and model."""
        self.api_key = api_key
        self.model = model
        self.base_url = "https://genos.genon.ai:3443"
        self.serving_id = 552
        
        # Setup headers
        self.headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.api_key}"
        }
    
    def chat_completion(self, prompt: str, max_tokens: int = 1000, temperature: float = 0):
        """Get a chat completion from GenOS."""
        try:
            # Prepare the request
            endpoint = f"{self.base_url}/api/gateway/rep/serving/{self.serving_id}/v1/chat/completions"
            
            data = {
                "model": self.model,
                "messages": [
                    {"role": "system", "content": "You are an expert at analyzing text and extracting information."},
                    {"role": "user", "content": prompt}
                ],
                "max_tokens": max_tokens,
                "temperature": temperature
            }
            
            # Make the API call with retry logic
            max_retries = 3
            retry_delay = 2  # seconds
            
            for attempt in range(max_retries):
                try:
                    response = requests.post(
                        endpoint,
                        headers=self.headers,
                        json=data,
                        timeout=60  # 60 second timeout
                    )
                    
                    # Check if response is successful
                    response.raise_for_status()
                    
                    # Handle both JSON and plain text responses
                    content_type = response.headers.get('Content-Type', '')
                    
                    if 'application/json' in content_type:
                        # JSON response
                        result = response.json()
                        if "choices" in result and len(result["choices"]) > 0:
                            return result["choices"][0]["message"]["content"]
                        return ""
                    else:
                        # Plain text response
                        logger.warning(f"Received non-JSON response: {response.text[:100]}...")
                        return response.text
                    
                except requests.exceptions.RequestException as e:
                    logger.warning(f"GenOS API request failed (attempt {attempt+1}/{max_retries}): {e}")
                    if attempt < max_retries - 1:
                        time.sleep(retry_delay)
                        retry_delay *= 2  # Exponential backoff
                    else:
                        logger.error(f"GenOS API request failed after {max_retries} attempts")
                        raise
                        
        except Exception as e:
            logger.error(f"Error calling GenOS API: {e}")
            return f"Error: {str(e)}"


class GenOSVectorMeta:
    """Class to represent vector metadata in GenOS format."""
    def __init__(self, **kwargs):
        self.text = kwargs.get('text')
        self.n_char = kwargs.get('n_char')
        self.n_word = kwargs.get('n_word')
        self.n_line = kwargs.get('n_line')
        self.i_page = kwargs.get('i_page')
        self.e_page = kwargs.get('e_page')
        self.i_chunk_on_page = kwargs.get('i_chunk_on_page')
        self.n_chunk_of_page = kwargs.get('n_chunk_of_page')
        self.i_chunk_on_doc = kwargs.get('i_chunk_on_doc')
        self.n_chunk_of_doc = kwargs.get('n_chunk_of_doc')
        self.n_page = kwargs.get('n_page')
        self.reg_date = kwargs.get('reg_date')
        self.chunk_bboxes = kwargs.get('chunk_bboxes')
        self.media_files = kwargs.get('media_files')
        
        # Additional fields for our date extraction
        self.meeting_date = kwargs.get('meeting_date')
        self.date_context = kwargs.get('date_context')
        self.all_dates = kwargs.get('all_dates')
        
    def model_validate(data):
        """Create a GenOSVectorMeta instance from a dictionary."""
        return GenOSVectorMeta(**data)
    
    def to_dict(self):
        """Convert to dictionary for serialization."""
        return {k: v for k, v in self.__dict__.items() if v is not None}


class DocumentProcessor:
    """Processor for documents with LLM-based date extraction."""
    
    def __init__(self):
        self.page_chunk_counts = {}
        self.llm = None  # Will be initialized when needed
        
    def extract_entities_with_llm(self, text: str) -> List[str]:
        """Extract entities from text using LLM."""
        # Initialize GenOS client if not already done
        if not hasattr(self, 'genos_client') or self.genos_client is None:
            self.genos_client = self.get_genos_client()
            
        if self.genos_client is None or not text.strip():
            return []
            
        # Create a prompt for entity extraction
        prompt = f"""
        Extract all key entities from the following text. Include people names, organizations, locations, dates, times, and any important terms or concepts.
        
        Return the entities as a JSON array of strings. Example:
        ```json
        ["John Smith", "Apple Inc.", "Seoul", "2023-05-15", "Marketing Department"]
        ```
        
        Text:
        {text[:1500]}  # Limit text length
        """
        
        try:
            # Get LLM response
            response = self.genos_client.chat_completion(prompt, max_tokens=500)
            
            # Try different patterns to extract entities
            # First try to find JSON array
            json_match = re.search(r'\[.*\]', response, re.DOTALL)
            if json_match:
                try:
                    entities_json = json_match.group()
                    entities = json.loads(entities_json)
                    if isinstance(entities, list):
                        return entities
                except json.JSONDecodeError:
                    pass
            
            # If no JSON array found, try to extract entities line by line
            entities = []
            for line in response.split('\n'):
                line = line.strip()
                # Remove common prefixes like "- ", "* ", "1. " etc.
                line = re.sub(r'^[-*•\d]+\.?\s*', '', line)
                # Remove quotes
                line = line.strip('"\'')
                if line and len(line) > 1 and not line.startswith(('```', '[')):
                    entities.append(line)
            
            return entities
        except Exception as e:
            logger.error(f"Error extracting entities with LLM: {e}")
            return []
            
    def generate_summary_with_llm(self, text: str) -> str:
        """Generate a summary of text using LLM."""
        # Initialize GenOS client if not already done
        if not hasattr(self, 'genos_client') or self.genos_client is None:
            self.genos_client = self.get_genos_client()
            
        if self.genos_client is None or not text.strip():
            return ""
            
        # Create a prompt for summary generation
        prompt = f"""
        Generate a brief summary (1-2 sentences) of the following text:
        
        {text[:1500]}  # Limit text length
        """
        
        try:
            # Get LLM response
            response = self.genos_client.chat_completion(prompt, max_tokens=100)
            return response.strip()
        except Exception as e:
            logger.error(f"Error generating summary with LLM: {e}")
            return ""
        
    def load_document_text(self, file_path: str) -> str:
        """Extract text from a document file."""
        ext = os.path.splitext(file_path)[-1].lower()
        if ext in ['.doc', '.docx']:
            # Use our own extraction method
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")
            
    def load_documents(self, file_path: str, **kwargs) -> List[Dict]:
        """Load documents and return as list of dictionaries with text and metadata."""
        text = self.load_document_text(file_path)
        
        # Create a simple document structure
        document = {
            "page_content": text,
            "metadata": {
                "source": file_path,
                "page": 1  # Default to page 1 for docx files
            }
        }
        
        return [document]
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file with metadata."""
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract metadata from document properties if available
            metadata = {}
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.created:
                    metadata["created"] = core_props.created.isoformat()
                if core_props.modified:
                    metadata["modified"] = core_props.modified.isoformat()
                if core_props.subject:
                    metadata["subject"] = core_props.subject
                if core_props.keywords:
                    metadata["keywords"] = core_props.keywords
            except Exception as e:
                logger.warning(f"Could not extract metadata from {file_path}: {e}")
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            # Add metadata as a header if available
            header_lines = []
            if "title" in metadata:
                header_lines.append(f"# {metadata['title']}")
            if "created" in metadata:
                created_date = metadata["created"].split("T")[0] if "T" in metadata["created"] else metadata["created"]
                header_lines.append(f"Date: {created_date}")
            if "author" in metadata:
                header_lines.append(f"Author: {metadata['author']}")
            if "subject" in metadata:
                header_lines.append(f"Subject: {metadata['subject']}")
            if "keywords" in metadata:
                header_lines.append(f"Keywords: {metadata['keywords']}")
            
            # Combine header and text
            if header_lines:
                full_text = header_lines + [""] + full_text
            
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX file {file_path}: {e}")
            return ""
    
    def extract_metadata_from_filename(self, file_path: str) -> Dict[str, str]:
        """Extract metadata from Korean meeting filename like '회의록_01_마케팅.docx'."""
        try:
            filename = os.path.basename(file_path)
            name_parts = filename.split('_')
            
            metadata = {
                "doc_type": "회의록" if len(name_parts) > 0 and name_parts[0] == "회의록" else "document",
            }
            
            # Extract meeting number if available
            if len(name_parts) > 1 and name_parts[1].isdigit():
                metadata["meeting_number"] = name_parts[1]
            
            # Extract meeting topic if available
            if len(name_parts) > 2:
                topic = name_parts[2].split('.')[0]  # Remove file extension
                metadata["topic"] = topic
            
            return metadata
        except Exception as e:
            logger.warning(f"Could not extract metadata from filename {file_path}: {e}")
            return {"doc_type": "document"}
    
    def get_genos_client(self):
        """Get a GenOS client for text processing."""
        try:
            return GenosClient(
                api_key="85a19cbab1454db29d0b1460304d679e",
                model="qwen2.5 72b instruct"
            )
        except Exception as e:
            logger.error(f"Error initializing GenOS client: {e}")
            return None
    
    def extract_dates_with_llm(self, text: str) -> Dict[str, Any]:
        """
        Extract dates from text using LLM.
        
        Args:
            text: The text to extract dates from
            
        Returns:
            Dictionary with extracted dates and their context
        """
        # Initialize GenOS client if not already done
        if not hasattr(self, 'genos_client') or self.genos_client is None:
            self.genos_client = self.get_genos_client()
            
        if self.genos_client is None:
            logger.error("Failed to initialize GenOS client")
            return {"primary_date": None, "all_dates": []}
            
        # Create the prompt for the LLM
        prompt = f"""
        You are an expert at analyzing text and extracting date information. 
        Extract ALL dates mentioned in the following text, paying special attention to meeting dates, event dates, and document dates.
        
        IMPORTANT RULES:
        1. Look for dates in various formats (YYYY-MM-DD, MM/DD/YYYY, Korean format like YYYY년 MM월 DD일, etc.)
        2. Pay special attention to phrases like "일시", "날짜", "Date:", "다음 회의:", etc. that indicate dates
        3. For each date found, identify its context (e.g., "meeting date", "document date", "next meeting", etc.)
        4. Always normalize dates to ISO format (YYYY-MM-DD)
        5. The PRIMARY date should be the actual meeting/event date, not the "next meeting" date
        6. Look for dates associated with "일시" (meaning "date/time") as these often indicate the actual meeting date
        
        Return a JSON object with the following structure:
        {{
          "primary_date": "YYYY-MM-DD",  // The actual meeting/event date (not "next meeting" date)
          "primary_date_context": "brief explanation of what this date represents",
          "all_dates": [
            {{
              "date": "YYYY-MM-DD",
              "context": "explanation of what this date represents",
              "original_text": "the original text snippet containing the date"
            }},
            ...
          ]
        }}
        
        Text:
        {text[:2000]}  // Truncate to first 2000 chars for LLM context window
        """
        
        try:
            # Make the API call
            response_text = self.genos_client.chat_completion(prompt, max_tokens=1500)
            
            # Parse the JSON response
            try:
                # Find JSON in the response
                json_match = re.search(r'```json\n(.*?)\n```', response_text, re.DOTALL)
                if json_match:
                    json_str = json_match.group(1)
                else:
                    # Try to find JSON without markdown formatting
                    json_match = re.search(r'({.*})', response_text, re.DOTALL)
                    if json_match:
                        json_str = json_match.group(1)
                    else:
                        json_str = response_text
                
                # Clean up the JSON string
                json_str = json_str.strip()
                if json_str.startswith('```') and json_str.endswith('```'):
                    json_str = json_str[3:-3].strip()
                    
                result = json.loads(json_str)
                return result
            except json.JSONDecodeError as e:
                logger.error(f"Error parsing LLM response: {e}")
                logger.debug(f"Response: {response_text}")
                return {"primary_date": None, "all_dates": []}
                
        except Exception as e:
            logger.error(f"Error extracting dates with LLM: {e}")
            return {"primary_date": None, "all_dates": []}
    
    def split_documents(self, documents: List[Dict], **kwargs) -> List[Dict]:
        """Split documents into chunks."""
        chunk_size = kwargs.get('chunk_size', 1000)
        chunk_overlap = kwargs.get('chunk_overlap', 200)
        
        chunks = []
        
        for document in documents:
            text = document.get("page_content", "")
            metadata = document.get("metadata", {})
            
            # Skip empty documents
            if not text.strip():
                continue
                
            # Simple text splitting by characters with overlap
            text_chunks = self._split_text(text, chunk_size, chunk_overlap)
            
            # Create chunks with metadata
            for i, chunk_text in enumerate(text_chunks):
                if not chunk_text.strip():
                    continue
                    
                chunk = {
                    "page_content": chunk_text,
                    "metadata": {
                        **metadata,
                        "chunk_index": i
                    }
                }
                chunks.append(chunk)
        
        if not chunks:
            raise Exception('Empty document')
        
        # Process page information
        self.page_chunk_counts = {}
        for chunk in chunks:
            page = chunk["metadata"].get('page', 1)
            
            # Ensure page is 1-based
            if isinstance(page, int) and page <= 0:
                page = 1
                
            chunk["metadata"]['page'] = page
            
            # Count chunks per page
            if page not in self.page_chunk_counts:
                self.page_chunk_counts[page] = 0
            self.page_chunk_counts[page] += 1
            
        return chunks
        
    def _split_text(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """Split text into chunks with overlap."""
        if not text:
            return []
            
        # Simple character-based chunking
        chunks = []
        start = 0
        text_len = len(text)
        
        while start < text_len:
            # Get chunk with specified size
            end = min(start + chunk_size, text_len)
            
            # Try to find a good breaking point (newline or space)
            if end < text_len:
                # Look for newline first
                newline_pos = text.rfind('\n', start, end)
                if newline_pos > start:
                    end = newline_pos + 1
                else:
                    # Look for space
                    space_pos = text.rfind(' ', start, end)
                    if space_pos > start:
                        end = space_pos + 1
            
            # Add the chunk
            chunks.append(text[start:end])
            
            # Move start position for next chunk, considering overlap
            start = max(start + chunk_size - chunk_overlap, end - chunk_overlap)
            
        return chunks
    
    def hybrid_chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Simplified hybrid chunking function that uses LLM to identify coherent chunks.
        
        Args:
            text: The text to chunk
            
        Returns:
            List of dictionaries with 'body', 'section_title', and other metadata
        """
        # Initialize GenOS client if not already done
        if not hasattr(self, 'genos_client') or self.genos_client is None:
            self.genos_client = self.get_genos_client()
            
        if self.genos_client is None:
            logger.error("Failed to initialize GenOS client")
            return self._fallback_chunking(text, [])
            
        # First, identify section headers to help guide the LLM
        section_headers = []
        section_pattern = r'(^|\n)(#+\s+.+)(\n|$)'
        for match in re.finditer(section_pattern, text):
            section_headers.append(match.group(2).strip())
        
        # Create the prompt for the LLM
        prompt = f"""
        You are an expert at analyzing and segmenting text. Split the following text into coherent, self-contained units of information.
        Each chunk should represent a single, complete unit of information that is understandable on its own.
        
        IMPORTANT RULES:
        1. NEVER split in the middle of sentences or paragraphs.
        2. Each section header (like "# Title", "## Section 2:", etc.) should start a new chunk.
        3. Make sure each chunk is a complete, coherent unit with proper context.
        4. Do not create chunks that start with conjunctions, partial sentences, or other fragments.
        5. Aim for chunks of around 200-300 words maximum.
        
        I've identified these section headers in the text:
        {section_headers}
        
        Return a JSON object with an array of chunks, where each chunk has:
        1. The complete text of the chunk (in the "text" field)
        2. The section title it belongs to (in the "section_title" field)
        3. A brief summary of what the chunk contains (in the "summary" field)
        
        Example response format:
        ```json
        {{
          "chunks": [
            {{
              "text": "Full text of the first chunk goes here...",
              "section_title": "Section 1: Introduction",
              "summary": "Brief description of what this chunk contains"
            }},
            ...
          ]
        }}
        ```
        
        Text to chunk:
        {text[:2000]}
        """
        
        try:
            # Get the LLM response
            response_text = self.genos_client.chat_completion(prompt, max_tokens=2000)
                
            # Find JSON in the response
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if not json_match:
                logger.warning("No JSON found in LLM response, falling back to simple chunking")
                return self._fallback_chunking(text, section_headers)
                
            result = json.loads(json_match.group())
            
            # Process the chunks
            chunks = []
            if "chunks" in result:
                for i, chunk_data in enumerate(result["chunks"]):
                    chunk_text = chunk_data.get("text", "").strip()
                    
                    # Skip empty chunks
                    if not chunk_text:
                        continue
                        
                    # Create the chunk object
                    chunk = {
                        "body": chunk_text,
                        "text": chunk_text,
                        "section_title": chunk_data.get("section_title", "Unknown Section"),
                        "summary": chunk_data.get("summary", ""),
                        "entities": [],
                        "relationships": []
                    }
                    chunks.append(chunk)
                    
                logger.info(f"LLM chunking produced {len(chunks)} chunks")
                return chunks
            else:
                logger.warning("No 'chunks' key found in LLM response, falling back to simple chunking")
                return self._fallback_chunking(text, section_headers)
                
        except Exception as e:
            logger.error(f"Error processing LLM chunks: {e}")
            return self._fallback_chunking(text, section_headers)
    
    def _fallback_chunking(self, text: str, section_headers: List[str]) -> List[Dict[str, Any]]:
        """
        Simple fallback chunking that splits by section headers.
        
        Args:
            text: The text to chunk
            section_headers: List of section headers found in the text
            
        Returns:
            List of dictionaries with 'body' and 'section_title'
        """
        logger.warning("Using fallback chunking method")
        
        # If no section headers, just return the whole text as one chunk
        if not section_headers:
            return [{
                "body": text.strip(),
                "text": text.strip(),
                "section_title": "Document",
                "summary": "Full document content",
                "entities": [],
                "relationships": []
            }]
        
        # Split by section headers
        chunks = []
        current_section = "Document"
        current_text = ""
        
        # Add the section headers as regex patterns
        patterns = [re.escape(header) for header in section_headers]
        section_pattern = r'(^|\n)(' + '|'.join(patterns) + r')(\n|$)'
        
        # Split by section headers
        parts = re.split(section_pattern, text)
        
        # Process the parts
        for i, part in enumerate(parts):
            if not part.strip():
                continue
                
            # Check if this is a section header
            if part.strip() in section_headers:
                # Save previous section
                if current_text.strip():
                    chunks.append({
                        "body": current_text.strip(),
                        "text": current_text.strip(),
                        "section_title": current_section,
                        "summary": f"Content from {current_section}",
                        "entities": [],
                        "relationships": []
                    })
                
                # Update current section
                current_section = part.strip()
                current_text = ""
            else:
                current_text += part
        
        # Add the final section
        if current_text.strip():
            chunks.append({
                "body": current_text.strip(),
                "text": current_text.strip(),
                "section_title": current_section,
                "summary": f"Content from {current_section}",
                "entities": [],
                "relationships": []
            })
        
        return chunks
    
    def compose_vectors(self, chunks: List[Dict], file_path: str, **kwargs) -> List[Dict]:
        """Compose vectors from chunks with metadata."""
        # Extract text for date extraction
        text = self.extract_text_from_docx(file_path)
        
        # Extract metadata from filename
        metadata = self.extract_metadata_from_filename(file_path)
        
        # Extract dates using LLM
        date_info = self.extract_dates_with_llm(text)
        meeting_date = date_info.get("primary_date")
        
        if meeting_date:
            # Convert YYYY-MM-DD to RFC3339 format
            try:
                # Parse the date and convert to RFC3339
                meeting_datetime = datetime.fromisoformat(meeting_date)
                meeting_date_rfc = meeting_datetime.isoformat(timespec='seconds') + 'Z'
                metadata["meeting_date"] = meeting_date_rfc
            except ValueError:
                # If parsing fails, keep the original format
                metadata["meeting_date"] = meeting_date
                
            metadata["date_context"] = date_info.get("primary_date_context", "")
            
            # Store all extracted dates in metadata
            all_dates = []
            for date_entry in date_info.get("all_dates", []):
                all_dates.append({
                    "date": date_entry.get("date"),
                    "context": date_entry.get("context"),
                    "original_text": date_entry.get("original_text")
                })
            metadata["all_dates"] = all_dates
            
            print(f"Extracted primary date: {meeting_date} ({metadata.get('date_context')}) from {file_path}")
            print(f"Found {len(all_dates)} dates in the document")
        
        # Get current timestamp in RFC3339 format
        current_time = datetime.now()
        rfc3339_timestamp = current_time.isoformat(timespec='seconds') + 'Z'
        
        # Global metadata for all chunks
        global_metadata = {
            "n_chunk_of_doc": len(chunks),
            "n_page": max([chunk["metadata"].get('page', 1) for chunk in chunks]),
            "reg_date": rfc3339_timestamp,  # RFC3339 format
            "created_at": rfc3339_timestamp  # RFC3339 format
        }
        
        current_page = None
        chunk_index_on_page = 0
        vectors = []
        
        # Process each chunk
        for chunk_idx, chunk in enumerate(chunks):
            page = chunk["metadata"].get('page', 1)
            text = chunk["page_content"]
            
            if page != current_page:
                current_page = page
                chunk_index_on_page = 0
            
            # Use our self-contained hybrid chunking to get better chunk quality
            hybrid_chunks = self.hybrid_chunk_text(text)
            
            # If hybrid chunking failed or returned empty, use the original chunk
            if not hybrid_chunks:
                # Generate a simple title from the first line
                first_line = ""
                if text:
                    first_line = text.split('\n', 1)[0].strip()
                    if len(first_line) > 100:  # Too long for a title
                        first_line = first_line[:97] + "..."
                
                # Extract entities using LLM
                llm_entities = self.extract_entities_with_llm(text)
                
                # Generate summary using LLM
                llm_summary = self.generate_summary_with_llm(text)
                
                hybrid_chunks = [{
                    "body": text,
                    "text": text,
                    "section_title": first_line or f"Section {chunk_idx + 1}",
                    "summary": llm_summary,
                    "entities": llm_entities,
                    "relationships": [],
                }]
            
            # Process each hybrid chunk
            for hybrid_chunk in hybrid_chunks:
                chunk_text = hybrid_chunk.get("body") or hybrid_chunk.get("text", "")
                
                # Skip empty chunks
                if not chunk_text:
                    continue
                
                # Extract entities, summary and title from the chunk
                chunk_entities = hybrid_chunk.get("entities", [])
                chunk_summary = hybrid_chunk.get("summary", "")
                
                # Determine title from section_title or first line
                chunk_title = hybrid_chunk.get("section_title", "")
                if not chunk_title and chunk_text:
                    # Use first line as title if no section title
                    first_line = chunk_text.split('\n', 1)[0].strip()
                    if len(first_line) > 5 and len(first_line) < 100:  # Reasonable title length
                        chunk_title = first_line
                
                # Create vector metadata with correct data types according to schema
                vector = GenOSVectorMeta(
                    # Text fields
                    text=chunk_text,
                    e_page=str(page),  # Text field in schema
                    meeting_date=metadata.get("meeting_date", ""),  # Text field (RFC3339 format)
                    date_context=metadata.get("date_context", ""),  # Text field
                    all_dates=json.dumps(metadata.get("all_dates", [])),  # Text field (JSON string)
                    media_files=json.dumps([]),  # Text field (JSON string)
                    chunk_bboxes=json.dumps([]),  # Text field (JSON string)
                    entities=json.dumps(chunk_entities),  # Add entities as JSON string
                    summary=chunk_summary,  # Add summary
                    title=chunk_title,  # Add title
                    
                    # Number fields
                    i_page=float(page),  # Number field (float) based on error
                    i_chunk_on_page=float(chunk_index_on_page),  # Number field (float) based on error
                    i_chunk_on_doc=float(chunk_idx),  # Number field (float) based on error
                    n_char=float(len(chunk_text)),  # Number field (ensure float)
                    n_word=float(len(chunk_text.split())),  # Number field (ensure float)
                    n_line=float(len(chunk_text.splitlines())),  # Number field (ensure float)
                    n_chunk_of_page=float(self.page_chunk_counts.get(page, 1)),  # Number field (ensure float)
                    n_chunk_of_doc=float(global_metadata["n_chunk_of_doc"]),  # Number field (ensure float)
                    n_page=float(global_metadata["n_page"]),  # Number field (ensure float)
                    
                    # Date fields
                    reg_date=global_metadata["reg_date"],  # Date field (ISO format)
                    created_at=global_metadata["created_at"]  # Integer timestamp in milliseconds
                )
                
                vectors.append(vector)
                chunk_index_on_page += 1
        
        return vectors
    
    async def __call__(self, request, file_path: str, **kwargs):
        """Process a DOCX file and return vectors."""
        # Load documents
        documents = self.load_documents(file_path, **kwargs)
        
        # Check if request is cancelled
        if hasattr(request, 'is_cancelled'):
            if await request.is_cancelled():
                return []
        
        # Split into chunks
        chunks = self.split_documents(documents, **kwargs)
        
        # Check if request is cancelled
        if hasattr(request, 'is_cancelled'):
            if await request.is_cancelled():
                return []
        
        # Compose vectors with metadata
        vectors = self.compose_vectors(chunks, file_path, **kwargs)
        
        # Check if request is cancelled
        if hasattr(request, 'is_cancelled'):
            if await request.is_cancelled():
                return []
        
        # Convert GenOSVectorMeta objects to dictionaries for serialization
        return [v.to_dict() for v in vectors]
