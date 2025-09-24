#!/usr/bin/env python3
"""
Clean reingestion script using the new single-pass LLM extraction system.
This replaces all the old ingestion scripts with a unified, modern approach.
"""

import os
import sys
import logging
import shutil
import json
import docx
from pathlib import Path
from typing import Dict, List, Any
from datetime import datetime

# Add project root to path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

from adapters.chroma_adapter import ChromaClient
# Import functions will be defined in this file

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/reingestion.log'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from DOCX including tables."""
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                full_text.append(' | '.join(row_text))
    
    return '\n'.join(full_text)

def llm_extract_all_metadata(text: str) -> dict:
    """Extract all metadata using a single LLM pass."""
    from configs.load import get_default_llm
    llm = get_default_llm()
    
    prompt = f"""
You are an expert at extracting key information from meeting minutes documents.
Analyze the following meeting minutes and extract the following details:
1. The actual meeting date (primary_date) in YYYY-MM-DD format.
2. A list of all attendees (attendees).
3. A list of key decisions (key_decisions).
4. A list of action items (action_items).
5. The meeting location (location).
6. The meeting topic (topic).

If a date is mentioned without a year, assume the current year (2025).
Return your answer as a JSON object with the following keys:
{{
  "primary_date": "YYYY-MM-DD",
  "attendees": ["Name1", "Name2"],
  "key_decisions": ["Decision 1", "Decision 2"],
  "action_items": ["Action 1", "Action 2"],
  "location": "Meeting Room Name",
  "topic": "Meeting Topic"
}}

Text:
{text}
"""
    
    try:
        response = llm.invoke(prompt)
        # Parse JSON from response
        import re
        json_match = re.search(r'\{.*\}', response.content, re.DOTALL)
        if json_match:
            metadata = json.loads(json_match.group())
            # Rename primary_date to meeting_date for consistency
            if 'primary_date' in metadata:
                metadata['meeting_date'] = metadata.pop('primary_date')
            return metadata
        else:
            logger.error("No JSON found in LLM response")
            return {}
    except Exception as e:
        logger.error(f"LLM extraction failed: {e}")
        return {}

def delete_chroma_db():
    """Delete the existing Chroma database for clean reingestion."""
    chroma_dir = project_root / "chroma_db"
    
    if chroma_dir.exists():
        logger.info(f"🗑️  Deleting existing Chroma database at {chroma_dir}")
        shutil.rmtree(chroma_dir)
        logger.info("✅ Chroma database deleted successfully")
        return True
    else:
        logger.info("ℹ️  Chroma database directory does not exist, nothing to delete")
        return True

def extract_text_from_docx_complete(file_path: str) -> str:
    """Extract all text from DOCX including tables (same as simple_ingestion.py)."""
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                full_text.append(' | '.join(row_text))
    
    return '\n'.join(full_text)

def reingest_document(file_path: str, client: ChromaClient) -> Dict[str, Any]:
    """
    Reingest a single document using the new single-pass LLM extraction system.
    
    Args:
        file_path: Path to the DOCX file
        client: ChromaDB client
        
    Returns:
        Dictionary with ingestion results and extracted metadata
    """
    logger.info(f"📄 Processing: {Path(file_path).name}")
    
    try:
        # Extract text from document
        text = extract_text_from_docx(file_path)
        logger.info(f"✅ Extracted {len(text)} characters")
        
        # Extract metadata with single LLM pass
        logger.info("🤖 Running single LLM pass for metadata extraction...")
        metadata = llm_extract_all_metadata(text)
        
        if not metadata:
            logger.error("❌ Failed to extract metadata from LLM")
            return {"error": "LLM extraction failed", "file": file_path}
        
        logger.info("✅ Metadata extracted successfully!")
        logger.info(f"📅 Meeting date: {metadata.get('meeting_date', 'NOT FOUND')}")
        logger.info(f"👥 Attendees: {len(metadata.get('attendees', []))}")
        logger.info(f"📋 Decisions: {len(metadata.get('key_decisions', []))}")
        logger.info(f"✅ Actions: {len(metadata.get('action_items', []))}")
        
        # Prepare document data for ChromaDB
        file_name = Path(file_path).stem
        doc_id = f"{file_name}_1"
        
        # Create chunks (for now, one chunk per document)
        chunk_data = {
            'chunk_id': f"{doc_id}:0",
            'doc_id': doc_id,
            'section': 'Document',
            'body': text,
            'entities': metadata.get('attendees', []),
            'valid_from': metadata.get('meeting_date', ''),
            'valid_to': '',
            'metadata': {
                **metadata,
                'file_name': file_name,
                'file_path': file_path,
                'ingestion_date': datetime.now().isoformat(),
                'valid_from': metadata.get('meeting_date', ''),
                'valid_to': '',
                'meeting_time': metadata.get('meeting_time', ''),
                'location': metadata.get('location', ''),
                'topic': metadata.get('topic', ''),
                'attendees': json.dumps(metadata.get('attendees', [])),
                'key_decisions': json.dumps(metadata.get('key_decisions', [])),
                'action_items': json.dumps(metadata.get('action_items', [])),
                'all_dates': json.dumps(metadata.get('all_dates', [])),
                'summary': metadata.get('summary', ''),
            }
        }
        
        # Ingest into ChromaDB
        collection = client._client.get_collection(client.chunk_collection)
        
        # Generate embeddings
        from configs.load import get_default_embeddings
        embeddings_model = get_default_embeddings()
        embeddings = embeddings_model.embed_documents([text])
        
        # Add to ChromaDB
        collection.add(
            documents=[text],
            metadatas=[chunk_data['metadata']],
            ids=[chunk_data['chunk_id']],
            embeddings=embeddings
        )
        
        logger.info(f"✅ Successfully ingested {file_name}")
        
        return {
            "success": True,
            "file": file_path,
            "doc_id": doc_id,
            "meeting_date": metadata.get('meeting_date'),
            "attendees": metadata.get('attendees', []),
            "decisions": len(metadata.get('key_decisions', [])),
            "actions": len(metadata.get('action_items', []))
        }
        
    except Exception as e:
        logger.error(f"❌ Error processing {file_path}: {e}")
        return {"error": str(e), "file": file_path}

def reingest_all_documents():
    """Reingest all documents in the data directory."""
    logger.info("🚀 Starting clean reingestion of all documents")
    
    # Delete existing database
    if not delete_chroma_db():
        logger.error("❌ Failed to delete existing database")
        return False
    
    # Connect to ChromaDB
    try:
        client = ChromaClient()
        logger.info("✅ Connected to ChromaDB")
    except Exception as e:
        logger.error(f"❌ Failed to connect to ChromaDB: {e}")
        return False
    
    # Get data directory
    data_dir = project_root / "data"
    
    if not data_dir.exists():
        logger.error(f"❌ Data directory not found: {data_dir}")
        return False
    
    # Find all DOCX files
    docx_files = list(data_dir.glob("*.docx"))
    
    if not docx_files:
        logger.error("❌ No DOCX files found in data directory")
        return False
    
    logger.info(f"📁 Found {len(docx_files)} DOCX files to process")
    
    # Process each document
    results = []
    successful = 0
    failed = 0
    
    for docx_file in sorted(docx_files):
        result = reingest_document(str(docx_file), client)
        results.append(result)
        
        if "error" in result:
            failed += 1
        else:
            successful += 1
    
    # Summary
    logger.info("=" * 60)
    logger.info("📊 REINGESTION SUMMARY")
    logger.info("=" * 60)
    logger.info(f"✅ Successful: {successful}")
    logger.info(f"❌ Failed: {failed}")
    logger.info(f"📁 Total files: {len(docx_files)}")
    
    # Show extracted dates
    logger.info("\n📅 EXTRACTED MEETING DATES:")
    for result in results:
        if "success" in result and result["success"]:
            file_name = Path(result["file"]).name
            meeting_date = result.get("meeting_date", "NOT FOUND")
            logger.info(f"  {file_name}: {meeting_date}")
    
    return successful > 0

def cleanup_old_files():
    """Delete old and irrelevant ingestion files."""
    logger.info("🧹 Cleaning up old and irrelevant files...")
    
    old_files = [
        "reingest_saibeo.py",
        "simple_ingestion.py", 
        "fix_saibeo_dates.py",
        "ingest_docx.py",
        "ingest_docx_llm.py",
        "ingest_beir_fiqa.py",
        "ingest_beir_fiqa_enhanced.py",
        "reset_and_reingest.py",
        "scripts/tools/reingest_all.py",
        "scripts/tools/check_ingested_dates.py",
        "scripts/tools/verify_ingestion.py",
    ]
    
    deleted_count = 0
    for file_path in old_files:
        full_path = project_root / file_path
        if full_path.exists():
            try:
                full_path.unlink()
                logger.info(f"🗑️  Deleted: {file_path}")
                deleted_count += 1
            except Exception as e:
                logger.warning(f"⚠️  Could not delete {file_path}: {e}")
    
    logger.info(f"✅ Cleaned up {deleted_count} old files")

def main():
    """Main function to run the clean reingestion."""
    logger.info("🎯 Starting clean reingestion process")
    
    # Clean up old files first
    cleanup_old_files()
    
    # Reingest all documents
    success = reingest_all_documents()
    
    if success:
        logger.info("🎉 Clean reingestion completed successfully!")
        logger.info("💡 You can now test the search functionality")
    else:
        logger.error("❌ Reingestion failed")
        sys.exit(1)

if __name__ == "__main__":
    main()
